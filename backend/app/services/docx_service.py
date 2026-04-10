import logging
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings

logger = logging.getLogger(__name__)

STORAGE_ROOT = Path(settings.STORAGE_PATH)


def generate_proposal(
    rfp_data: dict,
    extraction_data: dict,
    match_data: dict = None,
    company_name: str = "Your Company",
    template_path: str = None,
) -> str:
    """Generate a Word proposal document from AI extraction data."""

    if template_path and (STORAGE_ROOT / template_path).exists():
        doc = Document(str(STORAGE_ROOT / template_path))
    else:
        doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Cover Page
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROPOSAL")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(extraction_data.get("client_name", rfp_data.get("client_name", "")))
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")
    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scope_para.add_run(extraction_data.get("scope_summary", rfp_data.get("title", "")))
    run.font.size = Pt(14)

    doc.add_paragraph("")
    doc.add_paragraph("")
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_para.add_run(f"Prepared by {company_name}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
    run.font.size = Pt(12)

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    if match_data and match_data.get("executive_summary_draft"):
        doc.add_paragraph(match_data["executive_summary_draft"])
    elif extraction_data.get("scope_summary"):
        doc.add_paragraph(extraction_data["scope_summary"])

    # Project Scope
    doc.add_heading("2. Project Scope", level=1)
    doc.add_paragraph(extraction_data.get("scope_summary", ""))

    if extraction_data.get("timeline"):
        doc.add_heading("Timeline", level=2)
        doc.add_paragraph(extraction_data["timeline"])

    if extraction_data.get("budget_range"):
        doc.add_heading("Budget", level=2)
        doc.add_paragraph(extraction_data["budget_range"])

    # Technical Approach
    doc.add_heading("3. Technical Approach", level=1)
    key_reqs = extraction_data.get("key_requirements", [])
    if key_reqs:
        doc.add_paragraph("Our approach addresses the following key requirements:")
        for req in key_reqs:
            doc.add_paragraph(req, style="List Bullet")

    # If capability match data exists, add AI-written sections
    if match_data and match_data.get("section_matches"):
        doc.add_heading("4. Capability Alignment", level=1)
        for section in match_data["section_matches"]:
            doc.add_heading(section.get("rfp_requirement", ""), level=2)
            if section.get("ai_written_response"):
                doc.add_paragraph(section["ai_written_response"])
            score = section.get("match_score", 0)
            doc.add_paragraph(f"Match Score: {score}%")

    # Compliance Matrix
    compliance = extraction_data.get("compliance_requirements", [])
    if compliance:
        doc.add_heading("5. Compliance Matrix", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Requirement"
        hdr[1].text = "Compliance Status"
        hdr[2].text = "Evidence"
        for req in compliance:
            row = table.add_row().cells
            row[0].text = req
            row[1].text = "Compliant"
            row[2].text = "See attached documentation"

    # Team & References
    doc.add_heading("6. Team Structure", level=1)
    doc.add_paragraph("Our team brings extensive experience relevant to this engagement.")

    doc.add_heading("7. References", level=1)
    doc.add_paragraph("Available upon request.")

    # Save
    rfp_id = rfp_data.get("id", "unknown")
    version = rfp_data.get("next_version", 1)
    rel_path = f"generated/{rfp_id}/proposal_v{version}.docx"
    full_path = STORAGE_ROOT / rel_path
    full_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(full_path))

    return rel_path
